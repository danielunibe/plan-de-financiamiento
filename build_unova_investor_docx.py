from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "PROPUESTA_FORMAL_INVERSION_UNOVA_GAMES_STUDIO_2026.md"
OUTPUT = ROOT / "PROPUESTA_FORMAL_INVERSION_UNOVA_GAMES_STUDIO_2026.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(85, 85, 85)
MUTED = RGBColor(110, 110, 110)
BLACK = RGBColor(0, 0, 0)
LIGHT_FILL = "F4F6F9"
TABLE_HEADER = "E8EEF5"
CALIBRI = "Calibri"

ACCENT_MAP = {
    "accion": "acción",
    "aplicacion": "aplicación",
    "aportacion": "aportación",
    "asesoria": "asesoría",
    "ano": "año",
    "automatizacion": "automatización",
    "basico": "básico",
    "busqueda": "búsqueda",
    "capacitacion": "capacitación",
    "camara": "cámara",
    "clinicas": "clínicas",
    "codigo": "código",
    "coleccion": "colección",
    "comercializacion": "comercialización",
    "conexion": "conexión",
    "consolidacion": "consolidación",
    "critico": "crítico",
    "decision": "decisión",
    "desempeno": "desempeño",
    "dias": "días",
    "diseno": "diseño",
    "documentacion": "documentación",
    "economico": "económico",
    "educacion": "educación",
    "ejecucion": "ejecución",
    "energia": "energía",
    "estrategica": "estratégica",
    "estrategico": "estratégico",
    "exito": "éxito",
    "expansion": "expansión",
    "facil": "fácil",
    "formalizacion": "formalización",
    "informacion": "información",
    "infraestructura": "infraestructura",
    "inversion": "inversión",
    "inversionistas": "inversionistas",
    "juridico": "jurídico",
    "linea": "línea",
    "mas": "más",
    "medica": "médica",
    "medico": "médico",
    "Mexico": "México",
    "minimo": "mínimo",
    "notaria": "notaría",
    "notarias": "notarías",
    "notararias": "notarías",
    "operacion": "operación",
    "opcion": "opción",
    "participacion": "participación",
    "percepcion": "percepción",
    "podra": "podrá",
    "podran": "podrán",
    "posesion": "posesión",
    "publica": "pública",
    "publico": "público",
    "preformalizacion": "preformalización",
    "presentacion": "presentación",
    "proteccion": "protección",
    "reduccion": "reducción",
    "relacion": "relación",
    "revision": "revisión",
    "seria": "sería",
    "sera": "será",
    "si": "sí",
    "solucion": "solución",
    "tambien": "también",
    "tecnica": "técnica",
    "tecnico": "técnico",
    "tecnologia": "tecnología",
    "tecnologica": "tecnológica",
    "ultimo": "último",
    "util": "útil",
    "utiles": "útiles",
    "validacion": "validación",
    "visualizacion": "visualización",
}


def preserve_case(original, replacement):
    if original.isupper():
        return replacement.upper()
    if original[:1].isupper():
        return replacement[:1].upper() + replacement[1:]
    return replacement


def polish_span(text):
    if "http" in text:
        before, url = text.split("http", 1)
        return polish_span(before) + "http" + url
    if ".md" in text:
        return text
    for source, target in ACCENT_MAP.items():
        text = re.sub(
            rf"\b{re.escape(source)}\b",
            lambda match, repl=target: preserve_case(match.group(0), repl),
            text,
            flags=re.IGNORECASE,
        )
    return text


def set_run_font(run, name=CALIBRI, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in):
    width_dxa = int(width_in * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total_dxa = int(sum(widths) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D7DBE2")


def set_paragraph_tokens(paragraph, before=0, after=8, line=1.333, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_inline_runs(paragraph, text, size=11, color=BLACK, italic=False):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if bold else part
        clean = clean.replace("`", "")
        clean = polish_span(clean)
        run = paragraph.add_run(clean)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    set_paragraph_tokens(p, before=0, after=8, line=1.333, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_inline_runs(p, text, size=11, color=BLACK)
    return p


def add_quote(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    set_paragraph_tokens(p, before=0, after=0, line=1.25)
    add_inline_runs(p, text, size=11, color=NAVY, italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 0
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_decimal_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def add_list_item(doc, text, numbered=False, num_id=None, number=None):
    p = doc.add_paragraph(style="Normal")
    set_paragraph_tokens(p, before=0, after=4, line=1.208)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    prefix = f"{number}. " if numbered and number is not None else ("• " if not numbered else "")
    if prefix:
        marker = p.add_run(prefix)
        set_run_font(marker, size=11, color=BLACK, bold=False)
    add_inline_runs(p, text, size=11, color=BLACK)


def add_heading(doc, text, level):
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Heading 3")
    p = doc.add_paragraph(style=style)
    if level == 1:
        set_paragraph_tokens(p, before=18, after=10, line=1.1)
    elif level == 2:
        set_paragraph_tokens(p, before=12, after=6, line=1.1)
    else:
        set_paragraph_tokens(p, before=8, after=4, line=1.1)
    add_inline_runs(p, text, size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE)


def clean_table_cells(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def table_widths(headers):
    cols = len(headers)
    header_join = " ".join(headers).lower()
    if cols == 2:
        return [2.1, 4.4]
    if cols == 3 and "monto" in header_join:
        return [1.65, 3.45, 1.4]
    if cols == 3 and "nivel" in header_join:
        return [2.2, 0.8, 3.5]
    if cols == 3:
        return [1.45, 2.55, 2.5]
    if cols == 4:
        return [1.55, 1.0, 3.1, 0.85]
    return [6.5 / cols] * cols


def add_markdown_table(doc, rows):
    parsed = [clean_table_cells(row) for row in rows if not re.match(r"^\s*\|?\s*:?-{3,}", row)]
    if not parsed:
        return
    headers = parsed[0]
    body = parsed[1:]
    widths = table_widths(headers)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)

    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, TABLE_HEADER)
        p = cell.paragraphs[0]
        set_paragraph_tokens(p, before=0, after=0, line=1.1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(p, text, size=9.5, color=NAVY)
        for run in p.runs:
            run.bold = True

    for row in body:
        cells = table.add_row().cells
        for idx, text in enumerate(row[: len(headers)]):
            cell = cells[idx]
            p = cell.paragraphs[0]
            set_paragraph_tokens(p, before=0, after=0, line=1.15)
            if re.match(r"^\$|^Alto$|^Medio$|^Bajo$|^\d", text):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, text, size=9.2, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def set_base_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = CALIBRI
    normal._element.rPr.rFonts.set(qn("w:ascii"), CALIBRI)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), CALIBRI)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = CALIBRI
        style._element.rPr.rFonts.set(qn("w:ascii"), CALIBRI)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), CALIBRI)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def set_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    set_paragraph_tokens(p, before=0, after=0, line=1.0)
    r = p.add_run("Unova Games Studio | Propuesta confidencial de inversion")
    set_run_font(r, size=8.5, color=MUTED)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    set_paragraph_tokens(p, before=0, after=0, line=1.0)
    r = p.add_run("Documento privado - 17 de junio de 2026")
    set_run_font(r, size=8.5, color=MUTED)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_cover(doc):
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_tokens(p, before=0, after=6, line=1.1)
    r = p.add_run("UNOVA GAMES STUDIO")
    set_run_font(r, size=13, color=GRAY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_tokens(p, before=0, after=6, line=1.05)
    r = p.add_run("Propuesta formal de inversion")
    set_run_font(r, size=25, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_tokens(p, before=0, after=18, line=1.1)
    r = p.add_run("Consolidacion de estudio, pilotos comerciales y capital operativo de 90 dias")
    set_run_font(r, size=13, color=GRAY)

    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [2.2, 4.3])
    set_table_borders(table)
    metadata = [
        ("Preparado por", "Daniel Alexis Aguilar Unibe"),
        ("Fecha", "17 de junio de 2026"),
        ("Solicitud principal", "$120,000 MXN / 90 dias"),
        ("Formato sugerido", "2 inversionistas, $60,000 MXN cada uno, en 3 pagos mensuales"),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        for cell in row.cells:
            set_cell_margins(cell, top=100, bottom=100)
        set_cell_shading(row.cells[0], LIGHT_FILL)
        p = row.cells[0].paragraphs[0]
        set_paragraph_tokens(p, before=0, after=0, line=1.15)
        add_inline_runs(p, label, size=10, color=NAVY)
        for run in p.runs:
            run.bold = True
        p = row.cells[1].paragraphs[0]
        set_paragraph_tokens(p, before=0, after=0, line=1.15)
        add_inline_runs(p, value, size=10.5, color=BLACK)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    add_quote(
        doc,
        "La inversion no financia una promesa vacia. Financia la consolidacion comercial de un ano de trabajo tecnico y creativo ya acumulado.",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_tokens(p, before=10, after=4, line=1.1)
    r = p.add_run("Documento confidencial. Borrador comercial privado sujeto a revision legal y financiera.")
    set_run_font(r, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()


def parse_markdown_into_doc(doc, text):
    lines = text.splitlines()
    start = 0
    first_rule_seen = False
    for idx, line in enumerate(lines):
        if line.strip() == "---":
            first_rule_seen = True
            start = idx + 1
            break
    if not first_rule_seen:
        start = 0

    table_rows = []
    pending_para = []
    current_num_id = None

    def flush_para():
        nonlocal pending_para
        if pending_para:
            add_body_paragraph(doc, " ".join(pending_para).strip())
            pending_para = []

    def reset_numbering():
        nonlocal current_num_id
        current_num_id = None

    def flush_table():
        nonlocal table_rows
        if table_rows:
            flush_para()
            add_markdown_table(doc, table_rows)
            table_rows = []
            reset_numbering()

    for raw in lines[start:]:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_table()
            flush_para()
            reset_numbering()
            continue
        if stripped == "---":
            flush_table()
            flush_para()
            reset_numbering()
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            flush_para()
            reset_numbering()
            table_rows.append(stripped)
            continue
        flush_table()

        if stripped.startswith("### "):
            flush_para()
            reset_numbering()
            add_heading(doc, stripped[4:].strip(), 2)
        elif stripped.startswith("## "):
            flush_para()
            reset_numbering()
            add_heading(doc, stripped[3:].strip(), 1)
        elif stripped.startswith("# "):
            continue
        elif stripped.startswith("> "):
            flush_para()
            reset_numbering()
            add_quote(doc, stripped[2:].strip())
        elif re.match(r"^\d+\.\s+", stripped):
            flush_para()
            match = re.match(r"^(\d+)\.\s+(.*)", stripped)
            add_list_item(doc, match.group(2), numbered=True, number=int(match.group(1)))
        elif stripped.startswith("- "):
            flush_para()
            reset_numbering()
            add_list_item(doc, stripped[2:].strip(), numbered=False)
        else:
            reset_numbering()
            pending_para.append(stripped)
    flush_table()
    flush_para()


def main():
    doc = Document()
    set_base_styles(doc)
    set_header_footer(doc.sections[0])
    add_cover(doc)
    parse_markdown_into_doc(doc, SOURCE.read_text(encoding="utf-8"))

    core_properties = doc.core_properties
    core_properties.title = "Propuesta formal de inversion - Unova Games Studio"
    core_properties.subject = "Solicitud privada de inversion para consolidacion de estudio y pilotos de 90 dias"
    core_properties.author = "Daniel Alexis Aguilar Unibe"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
